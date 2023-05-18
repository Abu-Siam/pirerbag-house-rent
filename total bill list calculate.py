from docx import Document
# import comtypes.client
import time



def BanglaToEnglish(string_number: object) -> object:
    numbers = {'১': '1', '২': '2', '৩': '3', '৪': '4', '৫': '5', '৬': '6', '৭': '7', '৮': '8', '৯': '9', '০': '0'}
    english = ''
    for c in str(string_number):
        # print(numbers[c])
        english = english + numbers[c]
    # print(bangla)
    return str(english)


# print("wtf",int(BanglaToEnglish("১২৩"))+1)
rent_list=[]

document: object = Document("D:\Code\pirerbag bill\pirerbag-house-rent\doc\\total bill list.doc")
table = document.tables[0]

document1 = Document("D:\Code\pirerbag bill\pirerbag-house-rent\doc\\6N.doc")
table1 = document1.tables[0]
table.cell(1, 1).text = str(BanglaToEnglish(table1.cell(19, 1).text))
table.cell(1, 2).text = str(BanglaToEnglish(table1.cell(12, 1).text))
rent_list.append(int(table.cell(1, 1).text))


document2: object = Document("D:\Code\pirerbag bill\pirerbag-house-rent\doc\\6S.doc")
table2 = document2.tables[0]
table.cell(2, 1).text = str(BanglaToEnglish(table2.cell(19, 1).text))
table.cell(2, 2).text = str(BanglaToEnglish(table2.cell(12, 1).text))
rent_list.append(int(table.cell(2, 1).text))

#table.cell(14,1).text=str(c)


document3=Document("D:\Code\pirerbag bill\pirerbag-house-rent\doc\\5N.doc")
table3=document3.tables[0]
table.cell(3, 1).text = str(BanglaToEnglish(table3.cell(19, 1).text))
table.cell(3, 2).text = str(BanglaToEnglish(table3.cell(12, 1).text))
rent_list.append(int(table.cell(3, 1).text))


document4=Document("D:\Code\pirerbag bill\pirerbag-house-rent\doc\\5S.doc")
table4=document4.tables[0]
table.cell(4, 1).text = str(BanglaToEnglish(table4.cell(19, 1).text))
table.cell(4, 2).text = str(BanglaToEnglish(table4.cell(12, 1).text))
rent_list.append(int(table.cell(4, 1).text))


document5=Document("D:\Code\pirerbag bill\pirerbag-house-rent\doc\\4N.doc")
table5=document5.tables[0]
table.cell(5, 1).text = str(BanglaToEnglish(table5.cell(19, 1).text))
table.cell(5, 2).text = str(BanglaToEnglish(table5.cell(12, 1).text))
rent_list.append(int(table.cell(5, 1).text))


document6=Document("D:\Code\pirerbag bill\pirerbag-house-rent\doc\\4S.doc")
table6=document6.tables[0]
table.cell(6, 1).text = str(BanglaToEnglish(table6.cell(19, 1).text))
table.cell(6, 2).text = str(BanglaToEnglish(table6.cell(12, 1).text))
rent_list.append(int(table.cell(6, 1).text))


document7=Document("D:\Code\pirerbag bill\pirerbag-house-rent\doc\\3N.doc")
table7=document7.tables[0]
table.cell(7, 1).text = str(BanglaToEnglish(table7.cell(19, 1).text))
table.cell(7, 2).text = str(BanglaToEnglish(table7.cell(12, 1).text))
rent_list.append(int(table.cell(7, 1).text))


document8=Document("D:\Code\pirerbag bill\pirerbag-house-rent\doc\\3S.doc")
table8=document8.tables[0]
table.cell(8, 1).text = str(BanglaToEnglish(table8.cell(19, 1).text))
table.cell(8, 2).text = str(BanglaToEnglish(table8.cell(12, 1).text))
rent_list.append(int(table.cell(8, 1).text))


document9=Document("D:\Code\pirerbag bill\pirerbag-house-rent\doc\\2N.doc")
table9=document9.tables[0]
table.cell(9, 1).text = str(BanglaToEnglish(table9.cell(19, 1).text))
table.cell(9, 2).text = str(BanglaToEnglish(table9.cell(12, 1).text))
rent_list.append(int(table.cell(9, 1).text))


document10=Document("D:\Code\pirerbag bill\pirerbag-house-rent\doc\\2S.doc")
table10=document10.tables[0]
table.cell(10, 1).text = str(BanglaToEnglish(table10.cell(19, 1).text))
table.cell(10, 2).text = str(BanglaToEnglish(table10.cell(12, 1).text))
rent_list.append(int(table.cell(10, 1).text))


document11=Document("D:\Code\pirerbag bill\pirerbag-house-rent\doc\\1N.doc")
table11=document11.tables[0]
table.cell(11, 1).text = str(BanglaToEnglish(table11.cell(19, 1).text))
table.cell(11, 2).text = str(BanglaToEnglish(table11.cell(12, 1).text))
rent_list.append(int(table.cell(11, 1).text))


document12=Document("D:\Code\pirerbag bill\pirerbag-house-rent\doc\\1S.doc")
table12=document12.tables[0]
table.cell(12, 1).text = str(BanglaToEnglish(table12.cell(19, 1).text))
table.cell(12, 2).text = str(BanglaToEnglish(table12.cell(12, 1).text))
rent_list.append(int(table.cell(12, 1).text))

rent_list.append(int(table.cell(13,1).text))#VERY IMPORTANT 1ST FLOOR OLD TENANT DUES.MAY ADD MORE IN FUTURE


final_sum=0
for x in range(len(rent_list)):
    final_sum=final_sum+rent_list[x]
table.cell(15,1).text=str(final_sum)


document.save("D:\Code\pirerbag bill\pirerbag-house-rent\doc\\total bill list.doc")

wdFormatPDF = 17
temp_doc_string = "D:\Code\pirerbag bill\pirerbag-house-rent\doc\\total bill list.doc"
temp_pdf_string = "D:\Code\pirerbag bill\pirerbag-house-rent\doc\\total bill list.pdf"

in_file = temp_doc_string
out_file = temp_pdf_string

# word = comtypes.client.CreateObject('Word.Application')
# # key point 1: make word visible before open a new document
# word.Visible = True
# time.sleep(3)
#
# doc = word.Documents.Open(in_file)  # open docx file 1
# doc.SaveAs(out_file, FileFormat=wdFormatPDF)  # conversion
# doc.Close()  # close docx file 1
# word.Visible = False
# word.Quit()
