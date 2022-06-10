from builtins import object

from typing import Any, Union

from docx import Document

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from deprecation import *
import datetime
import pandas

import comtypes.client
import time
# english to bangla number translation


def EnglishToBangla(string_number : object) -> object:
    """

    :rtype: object
    """
    numbers = {'1': '১', '2': '২', '3': '৩', '4': '৪', '5': '৫', '6': '৬', '7': '৭', '8': '৮', '9': '৯', '0': '০'}
    bangla = ''
    for c in str(string_number):
        # print(numbers[c])
        bangla = bangla + numbers[c]
    # print(bangla)
    return str(bangla)


# doc file introduced


# GLOBAL VARIABLES!!!! MAY BE CHANGED FOR VARIOUS REASONS
rent_increase = 0
house_rent_list = {"5N": 10500, "4S": 11000, "4N": 11500, "3N": 6000, "3S": 10000, "2S": 11000, "1N": 7000, "5S": 10500,
                   "6N": 5000, "2N": 10400, "1S": 8000, "6S": 10000}
# prev_total_bill_list = {"6N": 0, "6S": 0, "5N": 0, "5S": 0, "4N": 0, "4S": 0, "3N": 0, "3S": 0, "2N": 0, "2S": 0,
#                       "1N": 0, "1S": 0}
# final_total_bill_list = {"6N": 0, "6S": 0, "5N": 0, "5S": 0, "4N": 0, "4S": 0, "3N": 0, "3S": 0, "2N": 0, "2S": 0,
#                       "1N": 0, "1S": 0}
now = datetime.datetime.now()

month_var = 0 # To CALCULATE PREVIOUS MONTH .ITS INTEGER EXPRESSES MONTH NO



estimated_current_bill: int = 500
while (1):
    checkmeter_bill = input("whats the checkmeter bill ? ")
    print(checkmeter_bill)
    verify = input("please verify again= ")
    if (verify == checkmeter_bill):
        break
    else:
        print("TYPE ERROR !!!!! !!!!try again")

gas_bill = 0


# calender = ["january", "february", "march", "april", "may", "june", "july", "august", "september", "october",
#           "november", "december"]

# FLAT NAME AND DATE IN THE TITLE
loop_flag = 0


while (1):
    document = Document("empty.doc")
    header: object = document.add_heading("বিল")
    header.style = 'Title'
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    while (1):
        flatno = input("whats the flat no ? ")
        print(flatno)
        verify = input("please verify again= ")
        if (verify == flatno):
            break
        else:
            print("TYPE ERROR !!!!! !!!!try again")

    temp_date = (pandas.Period(datetime.datetime.now(), 'M')  - month_var).strftime('%B \nYEAR : %Y')
    if (str(flatno).upper() == "6N"):
        water_bill = 0  # VERY VERY IMPORTANT MAY CHANGE IN FUTURE
        print("6 north floor water bill is reduced to 0")
    else : water_bill=500


    floorno = flatno[0]
    floorside = flatno[1]
    if (floorside == 'n'):
        floorside = 'north'
    elif (floorside == 's'):
        floorside = 'south'
    else:
        print("TYPE ERROR!!!!!!!!!")
   # title = "FLOOR : " + floorno + " (" + floorside + ") \t   DATE: " + str(now.day) + "-" + str(now.month-month_var) + "-" + str(
    #    now.year)
    title = "FLOOR : " + floorno + " (" + floorside + ") \t   MONTH: " + str(temp_date)
    title = title.upper()
    paragraph = document.add_heading(title, level=1)
    paragraph.bold = True
    paragraph.style = 'Title'
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # CREATING TABLE

    table = document.add_table(rows=25, cols=2)
    table.style = 'Table Grid'
    table.cell(19, 1).text = "01"  # for avoiding garbage value when making total bill list
    table.cell(0, 0).text = "বিষয়"
    table.cell(0, 1).text = "টাকা"
    if False:
        # EMPTY FLAT SITUATION HANDLE
        flat_empty_flag = input("is this flat empty ? (y/n)")
        if (flat_empty_flag == "y"):
            table.cell(19, 1).text = "০"

            temp = str(flatno).upper() + ".doc"
            document.save("F:\CODE\pirerbag bill\doc\\" + temp)
            break
        elif (flat_empty_flag == "n"):
            pass
        else:
            print("type error try again")
            break

        # NEW TENANT SITUATION HANDLE
        while (1):
            new_tenant_flag = input("is this new tenant? (y/n) ")
            if (new_tenant_flag == 'y'):
                print("new tenant")
                while (1):
                    advance_new_tenant = input("how much did they pay in advance ? ")
                    print(advance_new_tenant)
                    verify = input("please verify again= ")
                    if (verify == advance_new_tenant):
                        break
                    else:
                        print("TYPE ERROR !!!!! !!!!try again")
                temp_date = (pandas.Period(datetime.datetime.now(), 'M') - 1 - month_var).strftime('%B %Y')
                temp = str(temp_date) + " এর ADVANCE পরিশোধ"
                table.cell(1, 0).text = temp
                table.cell(1, 1).text = EnglishToBangla(advance_new_tenant)

                temp = str(temp_date) + " এর আনুমানিক বিদ্যুত বিল ধরা হল"
                table.cell(6, 0).text = temp
                table.cell(6, 1).text = EnglishToBangla(estimated_current_bill)

                temp = str(temp_date) + " এর গ্যাস বিল ধরা হল"
                table.cell(9, 0).text = temp
                table.cell(9, 1).text = EnglishToBangla(gas_bill)

                temp = str(temp_date) + " এর পানির বিল ধরা হল"
                table.cell(12, 0).text = temp
                table.cell(12, 1).text = EnglishToBangla(water_bill)

                house_rent = house_rent_list[str(flatno).upper()]
                print("house rent = ", house_rent)

                temp_date = (pandas.Period(datetime.datetime.now(), 'M') - month_var).strftime('%B %Y')
                temp = str(temp_date) + " এর বাসা ভাড়া"
                table.cell(15, 0).text = temp
                table.cell(15, 1).text = EnglishToBangla(house_rent)

                new_tenant_due_advance = int(house_rent) - int(advance_new_tenant)
                temp = "ADVANCE এর বকেয়া"
                table.cell(17, 0).text = temp
                table.cell(17, 1).text = EnglishToBangla(new_tenant_due_advance)

                new_tenant_bill = int(estimated_current_bill) + int(gas_bill) + int(water_bill) + int(house_rent) + int(
                    new_tenant_due_advance)
                print("hello", new_tenant_bill)

                temp = "এই মাসের( " + str(temp_date) + " ) এর total বিল (   " + str(
                    EnglishToBangla(estimated_current_bill)) + " + " + str(EnglishToBangla(gas_bill)) + " + " + str(
                    EnglishToBangla(water_bill)) + " + " + str(EnglishToBangla(house_rent)) + " + " + str(
                    EnglishToBangla(new_tenant_due_advance)) + " )"

                table.cell(19, 0).text = temp
                table.cell(19, 1).text = EnglishToBangla(new_tenant_bill)

                # EXTRA

                table.cell(21, 0).text = "বিবিধ"
                table.cell(21, 1).text = "-"
                table.cell(24, 0).text = "বাড়ীওয়ালার স্বাক্ষর"

                # SAVING IN DOC FILE

                temp = str(flatno).upper() + ".doc"
                document.save("F:\CODE\pirerbag bill\doc\\" + temp)

                break


            elif (new_tenant_flag == 'n'):
                print("old tenant")
                break
            else:
                print("TYPE ERROR !!!!! !!!!try again")
        if (new_tenant_flag == 'y'):
            while (1):
                confirmation = input("do you wanna try again? (y/n) ")
                #        print(EnglishToBangla(prev_desco_bill))
                #        verify = input("please verify again= ")
                if (confirmation == 'y'):
                    break
                elif (confirmation == 'n'):
                    loop_flag = 1
                    break
                else:
                    print("TYPE ERROR !!!!! !!!!try again")
            if (confirmation == 'n'):
                break
            else:
                continue

        print("program end checking")

# DESCO BILL CALCULATION

    temp_date = (pandas.Period(datetime.datetime.now(), 'M') - 2 - month_var).strftime(
        '%B %Y')  # AUTOMATIC MONTH CALCULATION
    temp: Union[str, Any] = str(temp_date) + " এর আনুমানিক বিদ্যুত বিল ধরা হয়েছিল"
    table.cell(1, 0).text = temp
    table.cell(1, 1).text = EnglishToBangla(estimated_current_bill)

    # GLOBAL VARIABLE(DESCO BILL) IS USED FOR CALCULATION FOR REDUCING COMPLEXITY


    while (1):
        prev_desco_bill = input("how much was actual desco bill? ")
        print(EnglishToBangla(prev_desco_bill))
        verify = input("please verify again= ")
        if (verify == prev_desco_bill):
            break
        else:
            print("TYPE ERROR !!!!! !!!!try again")

    temp = str(temp_date) + " এর ডেসকো বিদ্যুত বিল দিয়েছে"
    table.cell(2, 0).text = temp
    table.cell(2, 1).text = EnglishToBangla(str(prev_desco_bill))

    sum_desco = int(prev_desco_bill) - int(estimated_current_bill)

    # POSITIVE AND NEGATIVE VALUE OF DUE DESCO BILL  HANDLE

    if (sum_desco >= 0):
        temp = "টাকা দিবেন ( " + EnglishToBangla(prev_desco_bill) + " - " + EnglishToBangla(
            estimated_current_bill) + " )"
        temp_text = EnglishToBangla(abs(sum_desco))
    else:
        temp = "টাকা পাবেন ( " + EnglishToBangla(prev_desco_bill) + " - " + EnglishToBangla(
            estimated_current_bill) + " )"
        temp_text: str = "-" + str(EnglishToBangla(abs(sum_desco)))
    table.cell(4, 0).text = temp
    table.cell(4, 1).text = EnglishToBangla(abs(sum_desco))

    print("this much money they will pay or take after paying estimated desco bill= ", EnglishToBangla(abs(sum_desco)))

    # FIXED COST FOUND FROM GLOBAL VARIABLES

    print("checkmeter bill= ", EnglishToBangla(checkmeter_bill))
    print("estimate current bill  ", EnglishToBangla(estimated_current_bill))
    print("gas bill ", EnglishToBangla(gas_bill))
    print("water bill  ", EnglishToBangla(water_bill))
    table.cell(6, 0).text = "checkmeter বিল"
    table.cell(6, 1).text = EnglishToBangla(checkmeter_bill)

    temp_date = (pandas.Period(datetime.datetime.now(), 'M') - 1 - month_var).strftime('%B %Y')
    temp = str(temp_date) + " এর আনুমানিক বিদ্যুত বিল ধরা হল"
    table.cell(7, 0).text = temp
    table.cell(7, 1).text = EnglishToBangla(estimated_current_bill)
    temp = str(temp_date) + " এর গ্যাস বিল "
    table.cell(8, 0).text = temp
    table.cell(8, 1).text = EnglishToBangla(gas_bill)
    temp = str(temp_date) + " এর পানির বিল "
    table.cell(9, 0).text = temp
    table.cell(9, 1).text = EnglishToBangla(water_bill)

    # HOUSERENT FROM INPUT LATER IT WILL BE COLLECT FROM DATASETS
    """
    while (1):
        house_rent: int = input("whats the houserent of " + flatno + ' ? ')
        print(EnglishToBangla(house_rent))
        verify = input("please verify again= ")
        if (verify == house_rent):
            break
        else:
            print("TYPE ERROR !!!!! !!!!try again")
    """

    house_rent = house_rent_list[str(flatno).upper()]
    print("house rent = ", house_rent)
    temp_date = (pandas.Period(datetime.datetime.now(), 'M') - month_var).strftime('%B %Y')
    temp = str(temp_date) + " এর বাসা ভাড়া"
    table.cell(10, 0).text = temp
    table.cell(10, 1).text = EnglishToBangla(house_rent)

    temp_sum = int(sum_desco) + int(checkmeter_bill) + int(estimated_current_bill) + int(gas_bill) + int(
        water_bill) + int(house_rent)

    temp = "যোগফল ( " + temp_text + " + " + str(EnglishToBangla(checkmeter_bill)) + " + " + str(
        EnglishToBangla(estimated_current_bill)) + " + " + str(EnglishToBangla(gas_bill)) + " + " + str(
        EnglishToBangla(water_bill)) + " + " + str(
        EnglishToBangla(house_rent)) + " )"

    table.cell(12, 0).text = temp
    table.cell(12, 1).text = EnglishToBangla(temp_sum)

    # CALCULATION OF DUE MONEY

    while (1):
        prev_total_bill = input("whats the prev bill? ")
        print(EnglishToBangla(prev_total_bill))
        verify = input("please verify again= ")
        if (verify == prev_total_bill):
            break
        else:
            print("TYPE ERROR !!!!! !!!!try again")

    temp_date = (pandas.Period(datetime.datetime.now(), 'M') - 1 - month_var).strftime('%B %Y')
    temp = "আগের মাসের( " + str(temp_date) + " ) এর total বিল"
    table.cell(14, 0).text = temp
    table.cell(14, 1).text = EnglishToBangla(prev_total_bill)

    while (1):
        prev_paid_bill = input("how much did they pay in prev month? ")
        print(EnglishToBangla(prev_paid_bill))
        verify = input("please verify again= ")
        if (verify == prev_paid_bill):
            break
        else:
            print("TYPE ERROR !!!!! !!!!try again")

    temp = "আগের মাসে ( " + str(temp_date) + " ) পরিশোধ করেছিলেন"
    table.cell(15, 0).text = temp
    table.cell(15, 1).text = EnglishToBangla(prev_paid_bill)

    due_money = int(prev_total_bill) - int(prev_paid_bill)

    # CONSIDERING POSITIVE AND NEGATIVE VALUE of DUE MONEY

    if (due_money >= 0):
        temp = "বকেয়া ( " + EnglishToBangla(prev_total_bill) + " - " + EnglishToBangla(prev_paid_bill) + " )"
        temp_due = EnglishToBangla(abs(due_money))
    else:
        temp = "পাওনা ( " + EnglishToBangla(prev_total_bill) + " - " + EnglishToBangla(prev_paid_bill) + " )"
        temp_due = "- " + EnglishToBangla(abs(due_money))
    print("due money", due_money)

    table.cell(17, 0).text = temp
    table.cell(17, 1).text = EnglishToBangla(abs(due_money))

    # CALCULATION OF TOTAL BILL

    total_bill = int(sum_desco) + int(checkmeter_bill) + int(estimated_current_bill) + int(gas_bill) + int(
        water_bill) + int(house_rent) + int(due_money)

    print("total bill= ", EnglishToBangla(total_bill))

    temp_date = (pandas.Period(datetime.datetime.now(), 'M') - month_var).strftime('%B %Y')

    temp = "এই মাসের( " + str(temp_date) + " ) এর total বিল (   " + temp_due + " + " + EnglishToBangla(temp_sum) + " )"

    table.cell(19, 0).text = temp
    table.cell(19, 1).text = EnglishToBangla(total_bill)

    # EXTRA

    table.cell(21, 0).text = "বিবিধ"
    #table.cell(21, 1).text = "জানুয়ারী ২০১৯ থেকে ৫০০ ৳ ভাড়া বাড়ানো হয়েছে"
    table.cell(24, 0).text = "বাড়ীওয়ালার স্বাক্ষর"

    # SAVING IN DOC FILE

    temp = str(flatno).upper() + ".doc"
    document.save("F:\CODE\pirerbag bill\doc\\" + temp)

#CONVERTING TO PDF

    wdFormatPDF = 17
    temp_doc_string="F:\CODE\pirerbag bill\doc\\"+str(flatno).upper()+".doc"
    temp_pdf_string = "F:\CODE\pirerbag bill\doc\\" + str(flatno).upper() + ".pdf"

    in_file = temp_doc_string
    out_file =temp_pdf_string

    word = comtypes.client.CreateObject('Word.Application')
    # key point 1: make word visible before open a new document
    word.Visible = True
    time.sleep(3)

    doc = word.Documents.Open(in_file)  # open docx file 1
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)  # conversion
    doc.Close()  # close docx file 1
    word.Visible = False
    word.Quit()


    #REPEATATION
"""
    while (1):
        confirmation = input("do you wanna try again? (y/n) ")
        #        print(EnglishToBangla(prev_desco_bill))
        #        verify = input("please verify again= ")
        if (confirmation == 'y'):
            break
        elif (confirmation == 'n'):
            loop_flag = 1
            break
        else:
            print("TYPE ERROR !!!!! !!!!try again")"""

# LIST OF BILLS OF ALL TENANTS AND SAVING THEM IN SEPARATE DOC
