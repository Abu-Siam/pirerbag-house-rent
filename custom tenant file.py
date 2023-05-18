from builtins import object

from typing import Any, Union

from docx import Document

#from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from deprecation import *
import datetime
import pandas
# import comtypes.client
import time

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# english to bangla number translation

#########################
notice = "জুন  মাস, ২০২৩ থেকে পানির বিল ২০০ টাকা বাড়ানো হবে "
#############################

def EnglishToBangla(string_number : object) -> object:
    """

    :rtype: object
    """
    numbers = {'1': '১', '2': '২', '3': '৩', '4': '৪', '5': '৫', '6': '৬', '7': '৭', '8': '৮', '9': '৯', '0': '০'}
    bangla = ''
    for c in str(string_number):
        # print(numbers[c])
        bangla = bangla + numbers[c]
    # print(bangla)
    return str(bangla)


# doc file introducedyui


# GLOBAL VARIABLES!!!! MAY BE CHANGED FOR VARIOUS REASONS
rent_increase = 0
house_rent_list = {"5N": 10500, "4S": 10500, "4N": 11500, "3N": 6000, "3S": 10500, "2S": 11000, "1N": 7000, "5S": 10000,
                   "6N": 6000, "2N": 11000, "1S": 8000, "6S": 10000}
# prev_total_bill_list = {"6N": 0, "6S": 0, "5N": 0, "5S": 0, "4N": 0, "4S": 0, "3N": 0, "3S": 0, "2N": 0, "2S": 0,
#                       "1N": 0, "1S": 0}
# final_total_bill_list = {"6N": 0, "6S": 0, "5N": 0, "5S": 0, "4N": 0, "4S": 0, "3N": 0, "3S": 0, "2N": 0, "2S": 0,
#                       "1N": 0, "1S": 0}
now = datetime.datetime.now()

month_var = 0 # To CALCULATE PREVIOUS MONTH .ITS INTEGER EXPRESSES MONTH NO



# estimated_current_bill: int = 500
# checkmeter_bill = input("whats the checkmeter bill ? ")
# print("Cheking Checkmeter bill ="+checkmeter_bill)


# gas_bill = 0


# calender = ["january", "february", "march", "april", "may", "june", "july", "august", "september", "october",
#           "november", "december"]

# FLAT NAME AND DATE IN THE TITLE
loop_flag = 0


while (1):
    
    document = Document("empty.doc")
    header: object = document.add_heading("বিল")
    header.style = 'Title'
    header.style.font.size = 25
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    flatno = input("whats the flat no ? ")
    print("Checking Flat no : " + flatno)


    temp_date = (pandas.Period(datetime.datetime.now(), 'M')  - month_var).strftime('%B\t YEAR : %Y')
    water_bill=500


    floorno = flatno[0]
    floorside = flatno[1]
    if (floorside == 'n'):
        floorside = 'north'
    elif (floorside == 's'):
        floorside = 'south'
    else:
        print("TYPE ERROR!!!!!!!!!")
   # title = "FLOOR : " + floorno + " (" + floorside + ") \t   DATE: " + str(now.day) + "-" + str(now.month-month_var) + "-" + str(
    #    now.year)
    title = "Floor: " + floorno + " (" + floorside + ")\tMonth: " + str(temp_date)

    paragraph = document.add_heading(title, level=2)
    # paragraph.bold = True
    paragraph.style = 'Title'
    paragraph.style.font.size = 23
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # CREATING TABLE

    table = document.add_table(rows=15, cols=2)
    table.style = 'Table Grid'
    table.cell(9, 1).text = "01"  # for avoiding garbage value when making total bill list
    table.cell(0, 0).text = "বিষয়"
    table.cell(0, 1).text = "টাকা"


# DESCO BILL CALCULATION

    temp_date = (pandas.Period(datetime.datetime.now(), 'M') - 2 - month_var).strftime(
        '%B %Y')  # AUTOMATIC MONTH CALCULATION
    # temp: Union[str, Any] = str(temp_date) + " এর আনুমানিক বিদ্যুত বিল ধরা হয়েছিল"
    # table.cell(1, 0).text = temp
    # table.cell(1, 1).text = EnglishToBangla(estimated_current_bill)

    # GLOBAL VARIABLE(DESCO BILL) IS USED FOR CALCULATION FOR REDUCING COMPLEXITY



    # prev_desco_bill = input("how much was actual desco bill? ")
    # print("Checkinng Desco Bill = " + EnglishToBangla(prev_desco_bill))
    #
    #
    # temp = str(temp_date) + " এর ডেসকো বিদ্যুত বিল দিয়েছে"
    # table.cell(2, 0).text = temp
    # table.cell(2, 1).text = EnglishToBangla(str(prev_desco_bill))

    # sum_desco = int(prev_desco_bill) - int(estimated_current_bill)

    # POSITIVE AND NEGATIVE VALUE OF DUE DESCO BILL  HANDLE

    # if (sum_desco >= 0):
    #     temp = "টাকা দিবেন ( " + EnglishToBangla(prev_desco_bill) + " - " + EnglishToBangla(
    #         estimated_current_bill) + " )"
    #     temp_text = EnglishToBangla(abs(sum_desco))
    # else:
    #     temp = "টাকা পাবেন ( " + EnglishToBangla(prev_desco_bill) + " - " + EnglishToBangla(
    #         estimated_current_bill) + " )"
    #     temp_text: str = "-" + str(EnglishToBangla(abs(sum_desco)))
    # table.cell(4, 0).text = temp
    # table.cell(4, 1).text = EnglishToBangla(abs(sum_desco))

    # print("this much money they will pay or take after paying estimated desco bill= ", EnglishToBangla(abs(sum_desco)))

    # FIXED COST FOUND FROM GLOBAL VARIABLES

    #print("checkmeter bill= ", EnglishToBangla(checkmeter_bill))
    #print("estimate current bill  ", EnglishToBangla(estimated_current_bill))
    #print("gas bill ", EnglishToBangla(gas_bill))
    #print("water bill  ", EnglishToBangla(water_bill))
    # table.cell(6, 0).text = "checkmeter & pump বিল"
    # table.cell(6, 1).text = EnglishToBangla(checkmeter_bill)

    temp_date = (pandas.Period(datetime.datetime.now(), 'M') - 1 - month_var).strftime('%B %Y')
    # temp = str(temp_date) + " এর আনুমানিক বিদ্যুত বিল ধরা হল"
    # table.cell(7, 0).text = temp
    # table.cell(7, 1).text = EnglishToBangla(estimated_current_bill)
    # temp = str(temp_date) + " এর গ্যাস বিল "
    # table.cell(8, 0).text = temp
    # table.cell(8, 1).text = EnglishToBangla(gas_bill)
    temp = str(temp_date) + " এর পানির বিল "
    table.cell(1, 0).text = temp
    table.cell(1, 1).text = EnglishToBangla(water_bill)



    house_rent = house_rent_list[str(flatno).upper()]
    print("house rent = ", house_rent)
    temp_date = (pandas.Period(datetime.datetime.now(), 'M') - month_var).strftime('%B %Y')
    temp = str(temp_date) + " এর বাসা ভাড়া"
    table.cell(2, 0).text = temp
    table.cell(2, 1).text = EnglishToBangla(house_rent)

    temp_sum = int(water_bill) + int(house_rent)

    temp = "যোগফল ( " + str(
        EnglishToBangla(water_bill)) + " + " + str(
        EnglishToBangla(house_rent)) + " )"

    table.cell(3, 0).text = temp
    table.cell(3, 1).text = EnglishToBangla(temp_sum)

    # CALCULATION OF DUE MONEY


    prev_total_bill = input("whats the prev bill? ")
    print("Checking previous month due = " + EnglishToBangla(prev_total_bill))


    temp_date = (pandas.Period(datetime.datetime.now(), 'M') - 1 - month_var).strftime('%B %Y')
    temp = "আগের মাসের( " + str(temp_date) + " ) এর total বিল"
    table.cell(4, 0).text = temp
    table.cell(4, 1).text = EnglishToBangla(prev_total_bill)


    prev_paid_bill = input("how much did they pay in prev month? ")
    print("Checking tk paid in previous month = " + EnglishToBangla(prev_paid_bill))


    temp = "আগের মাসে ( " + str(temp_date) + " ) পরিশোধ করেছিলেন"
    table.cell(5, 0).text = temp
    table.cell(5, 1).text = EnglishToBangla(prev_paid_bill)

    due_money = int(prev_total_bill) - int(prev_paid_bill)

    # CONSIDERING POSITIVE AND NEGATIVE VALUE of DUE MONEY

    if (due_money >= 0):
        temp = "বকেয়া ( " + EnglishToBangla(prev_total_bill) + " - " + EnglishToBangla(prev_paid_bill) + " )"
        temp_due = EnglishToBangla(abs(due_money))
    else:
        temp = "পাওনা ( " + EnglishToBangla(prev_total_bill) + " - " + EnglishToBangla(prev_paid_bill) + " )"
        temp_due = "- " + EnglishToBangla(abs(due_money))
    print("due money", due_money)
    if(due_money < 10 and due_money > 0):
        due_money = 0
    table.cell(7, 0).text = temp
    table.cell(7, 1).text = EnglishToBangla(abs(due_money))

    # CALCULATION OF TOTAL BILL

    total_bill =  int(water_bill) + int(house_rent) + int(due_money)

    print("total bill= ", EnglishToBangla(total_bill))

    temp_date = (pandas.Period(datetime.datetime.now(), 'M') - month_var).strftime('%B %Y')

    temp = "এই মাসের( " + str(temp_date) + " ) এর total বিল (   " + temp_due + " + " + EnglishToBangla(temp_sum) + " )"

    table.cell(9, 0).text = temp
    table.cell(9, 1).text = EnglishToBangla(total_bill)

    # EXTRA

    table.cell(11, 0).text = "বিবিধ"
    table.cell(11, 1).text = notice
    #table.cell(21, 1).text = "জানুয়ারী ২০১৯ থেকে ৫০০ ৳ ভাড়া বাড়ানো হয়েছে"
    table.cell(12, 0).text = "বাড়ীওয়ালার স্বাক্ষর"
    table.cell(12, 1).text = 'Abu Siam(mob: 01938513816),\n\nSultana Razia Zummi(mob: 01674069844),\n\nAbu Zakaria(zakariatanim@gmail.com)'
    table.cell(12, 1).style = 'List Bullet'
    table.cell(13, 0).text = "স্থানীয় কর্তৃপক্ষ"
    table.cell(13, 1).text = "Zahirul Islam Chowdhury (mob:01676948276)"




    # SAVING IN DOC FILE

    temp = str(flatno).upper() + ".doc"
    document.save("./doc/" + temp)

#CONVERTING TO PDF

    # wdFormatPDF = 17
    # temp_doc_string="F:\CODE\pirerbag bill\doc\\"+str(flatno).upper()+".doc"
    # temp_pdf_string = "F:\CODE\pirerbag bill\doc\\" + str(flatno).upper() + ".pdf"
    #
    # in_file = temp_doc_string
    # out_file =temp_pdf_string
    #
    # word = comtypes.client.CreateObject('Word.Application')
    # # key point 1: make word visible before open a new document
    # word.Visible = True
    # time.sleep(3)
    #
    # doc = word.Documents.Open(in_file)  # open docx file 1
    # doc.SaveAs(out_file, FileFormat=wdFormatPDF)  # conversion
    # doc.Close()  # close docx file 1
    # word.Visible = False
    # word.Quit()


    #REPEATATION
"""
    while (1):
        confirmation = input("do you wanna try again? (y/n) ")
        #        print(EnglishToBangla(prev_desco_bill))
        #        verify = input("please verify again= ")
        if (confirmation == 'y'):
            break
        elif (confirmation == 'n'):
            loop_flag = 1
            break
        else:
            print("TYPE ERROR !!!!! !!!!try again")"""

# LIST OF BILLS OF ALL TENANTS AND SAVING THEM IN SEPARATE DOC
