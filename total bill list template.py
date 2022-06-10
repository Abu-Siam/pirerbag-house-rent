import time
#u need to import python-docx
import comtypes.client

from docx import Document



import datetime

now = datetime.datetime.now()
document2: object = Document("empty.doc")
title2 = "TOTAL HOUSE বিল  " + "\t  DATE: " + str(now.day) + "-" + str(now.month) + "-" + str(now.year)
# title2 = title2.upper()
paragraph2 = document2.add_heading(title2, level=1)
paragraph2.bold = True
paragraph2.style = 'Title'
#paragraph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
table2: object = document2.add_table(rows=19, cols=3)
table2.style = 'Table Grid'
table2.cell(0, 0).text = "বাসা"
table2.cell(0, 1).text = "total bill"
table2.cell(0, 2).text = "এই মাসের বিল"
table2.cell(1, 0).text = "6th north"
table2.cell(1, 1).text = "0"
table2.cell(1, 2).text = "0"
table2.cell(2, 0).text = "6th south"
table2.cell(2, 1).text = "0"
table2.cell(2, 2).text = "0"
table2.cell(3, 0).text = "5th north"
table2.cell(3, 1).text = "0"
table2.cell(3, 2).text = "0"
table2.cell(4, 0).text = "5th south"
table2.cell(4, 1).text = "0"
table2.cell(4, 2).text = "0"
table2.cell(5, 0).text = "4th north"
table2.cell(5, 1).text = "0"
table2.cell(5, 2).text = "0"
table2.cell(6, 0).text = "4th south"
table2.cell(6, 1).text = "0"
table2.cell(6, 2).text = "0"
table2.cell(7, 0).text = "3rd north"
table2.cell(7, 1).text = "0"
table2.cell(7, 2).text = "0"
table2.cell(8, 0).text = "3rd south"
table2.cell(8, 1).text = "0"
table2.cell(8, 2).text = "0"
table2.cell(9, 0).text = "2nd north"
table2.cell(9, 1).text = "0"
table2.cell(9, 2).text = "0"
table2.cell(10, 0).text = "2nd south"
table2.cell(10, 1).text = "0"
table2.cell(10, 2).text = "0"
table2.cell(11, 0).text = "1st north"
table2.cell(11, 1).text = "0"
table2.cell(11, 2).text = "0"
table2.cell(12, 0).text = "1st south"
table2.cell(12, 1).text = "0"
table2.cell(12, 2).text = "0"
table2.cell(13, 0).text = "1st floor(north)(old)"
table2.cell(13, 1).text = '18331'
table2.cell(15, 0).text = "total (যোগফল) ="
table2.cell(18, 0).text = "বাড়ীওয়ালার স্বাক্ষর"
table2.cell(18, 1).text = ""

document2.save("F:\CODE\pirerbag bill\doc\\total bill list.doc")

"""
wdFormatPDF = 17
temp_doc_string = "F:\CODE\pirerbag bill\doc\\total bill list.doc"
temp_pdf_string = "F:\CODE\pirerbag bill\doc\\total bill list.pdf"

in_file = temp_doc_string
out_file = temp_pdf_string

word = comtypes.client.CreateObject('Word.Application')
# key point 1: make word visible before open a new document
word.Visible = True
#docx.time.sleep(3)

doc = word.Documents.Open(in_file)  # open docx file 1
doc.SaveAs(out_file, FileFormat=wdFormatPDF)  # conversion
doc.Close()  # close docx file 1
word.Visible = False
word.Quit()"""


document3: object = Document("empty.doc")
title3 = "TOTAL APARTMENT BILL  " + "\t  DATE: " + str(now.day) + "-" + str(now.month) + "-" + str(now.year)
# title2 = title2.upper()
paragraph3 = document3.add_heading(title3, level=1)
paragraph3.bold = True
paragraph3.style = 'Title'

table3: object = document3.add_table(rows=8, cols=2)
table3.style = 'Table Grid'
table3.cell(0, 0).text = "বাসা"
table3.cell(0, 1).text = "Amount"

table3.cell(1, 0).text = "A1 apartment rent"
table3.cell(1, 1).text = "27000"

table3.cell(2, 0).text = "gas bill"
table3.cell(2, 1).text = "975"

table3.cell(3, 0).text = "due "

apt_rent_due:int=input("whats the due of a1=")
table3.cell(3, 1).text = str(apt_rent_due)

table3.cell(4, 0).text = "extra"
extra: int=0
table3.cell(4, 1).text = str(extra)

table3.cell(5, 0).text = "total bill"

sum3 =27000+975+int(apt_rent_due)+extra
table3.cell(5, 1).text = str(sum3)

table3.cell(6, 0).text = ""
table3.cell(6, 1).text = ""

table3.cell(7, 0).text = "signature of owner"
table3.cell(7, 1).text = ""


document3.save("F:\CODE\pirerbag bill\doc\\bill A1 apartment.doc")
wdFormatPDF = 17
temp_doc_string = "F:\CODE\pirerbag bill\doc\\bill A1 apartment.doc"
temp_pdf_string = "F:\CODE\pirerbag bill\doc\\bill A1 apartment.pdf"

in_file = temp_doc_string
out_file = temp_pdf_string

# word = comtypes.client.CreateObject('Word.Application')
# # key point 1: make word visible before open a new document
# word.Visible = True
# #docx.time.sleep(3)
#
# doc = word.Documents.Open(in_file)  # open docx file 1
# doc.SaveAs(out_file, FileFormat=wdFormatPDF)  # conversion
# doc.Close()  # close docx file 1
# word.Visible = False
# word.Quit()

