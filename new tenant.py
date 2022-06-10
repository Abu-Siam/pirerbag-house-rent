from typing import Any, Union

from docx import Document

from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import pandas
import os
import comtypes.client
import time
# english to bangla number translation


def EnglishToBangla(string_number: object) -> object:
    """

    :rtype: object
    """
    numbers = {'1': '১', '2': '২', '3': '৩', '4': '৪', '5': '৫', '6': '৬', '7': '৭', '8': '৮', '9': '৯', '0': '০'}
    bangla = ''
    for c in str(string_number):
        # print(numbers[c])
        bangla = bangla + numbers[c]
    # print(bangla)
    return str(bangla)


# doc file introduced


# GLOBAL VARIABLES!!!! MAY BE CHANGED FOR VARIOUS REASONS
rent_increase = 0
house_rent_list = {"5N": 10500, "4S": 11000, "4N": 11500, "3N": 6000, "3S": 10500, "2S": 11000, "1N": 7000, "5S": 10500,
                   "6N": 5000, "2N": 11000, "1S": 9500, "6S": 10000}
# prev_total_bill_list = {"6N": 0, "6S": 0, "5N": 0, "5S": 0, "4N": 0, "4S": 0, "3N": 0, "3S": 0, "2N": 0, "2S": 0,
#                       "1N": 0, "1S": 0}
# final_total_bill_list = {"6N": 0, "6S": 0, "5N": 0, "5S": 0, "4N": 0, "4S": 0, "3N": 0, "3S": 0, "2N": 0, "2S": 0,
#                       "1N": 0, "1S": 0}
now = datetime.datetime.now()

month_var = 0  # To CALCULATE PREVIOUS MONTH .ITS INTEGER EXPRESSES MONTH NO


estimated_current_bill: int = 500
while (1):
    checkmeter_bill = input("whats the checkmeter bill ? ")
    print(checkmeter_bill)
    verify = input("please verify again= ")
    if (verify == checkmeter_bill):
        break
    else:
        print("TYPE ERROR !!!!! !!!!try again")

gas_bill = 800
water_bill=500
document = Document()
header: object = document.add_heading("বিল")
header.style = 'Title'
header.alignment = WD_ALIGN_PARAGRAPH.CENTER

while (1):
    flatno = input("whats the flat no ? ")
    print(flatno)
    verify = input("please verify again= ")
    if (verify == flatno):
        break
    else:
        print("TYPE ERROR !!!!! !!!!try again")


floorno = flatno[0]
floorside = flatno[1]
if (floorside == 'n'):
    floorside = 'north'
elif (floorside == 's'):
    floorside = 'south'
else:
    print("TYPE ERROR!!!!!!!!!")
title = "FLOOR : " + floorno + " (" + floorside + ") \t   DATE: " + str(now.day) + "-" + str(now.month) + "-" + str(
    now.year)
title = title.upper()
paragraph = document.add_heading(title, level=1)
paragraph.bold = True
paragraph.style = 'Title'
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# CREATING TABLE

table = document.add_table(rows=25, cols=2)
table.style = 'Table Grid'
table.cell(19, 1).text = "01"  # for avoiding garbage value when making total bill list
table.cell(0, 0).text = "বিষয়"
table.cell(0, 1).text = "টাকা"

print("new tenant")
while (1):
    advance_new_tenant = input("how much did they pay in advance ? ")
    print(advance_new_tenant)
    verify = input("please verify again= ")
    if (verify == advance_new_tenant):
        break
    else:
        print("TYPE ERROR !!!!! !!!!try again")
temp_date = (pandas.Period(datetime.datetime.now(), 'M') - 1 - month_var).strftime('%B %Y')
temp = str(temp_date) + " এর ADVANCE পরিশোধ"
table.cell(1, 0).text = temp
table.cell(1, 1).text = EnglishToBangla(advance_new_tenant)

temp = str(temp_date) + " এর আনুমানিক বিদ্যুত বিল ধরা হল"
table.cell(6, 0).text = temp
table.cell(6, 1).text = EnglishToBangla(estimated_current_bill)

temp = str(temp_date) + " এর গ্যাস বিল ধরা হল"
table.cell(9, 0).text = temp
table.cell(9, 1).text = EnglishToBangla(gas_bill)

temp = str(temp_date) + " এর পানির বিল ধরা হল"
table.cell(12, 0).text = temp
table.cell(12, 1).text = EnglishToBangla(water_bill)

house_rent = house_rent_list[str(flatno).upper()]
print("house rent = ", house_rent)

temp_date = (pandas.Period(datetime.datetime.now(), 'M') - month_var).strftime('%B %Y')
temp = str(temp_date) + " এর বাসা ভাড়া"
table.cell(15, 0).text = temp
table.cell(15, 1).text = EnglishToBangla(house_rent)

new_tenant_due_advance = int(house_rent) - int(advance_new_tenant)
temp = "ADVANCE এর বকেয়া"
table.cell(17, 0).text = temp
table.cell(17, 1).text = EnglishToBangla(new_tenant_due_advance)

new_tenant_bill = int(estimated_current_bill) + int(gas_bill) + int(water_bill) + int(house_rent) + int(
    new_tenant_due_advance)
print("hello", new_tenant_bill)

temp = "এই মাসের( " + str(temp_date) + " ) এর total বিল (   " + str(
    EnglishToBangla(estimated_current_bill)) + " + " + str(EnglishToBangla(gas_bill)) + " + " + str(
    EnglishToBangla(water_bill)) + " + " + str(EnglishToBangla(house_rent)) + " + " + str(
    EnglishToBangla(new_tenant_due_advance)) + " )"

table.cell(19, 0).text = temp
table.cell(19, 1).text = EnglishToBangla(new_tenant_bill)

# EXTRA

table.cell(21, 0).text = "বিবিধ"
table.cell(21, 1).text = "-"
table.cell(24, 0).text = "বাড়ীওয়ালার স্বাক্ষর"

# SAVING IN DOC FILE

temp = str(flatno).upper() + ".doc"
document.save("F:\CODE\pirerbag bill\doc\\" + temp)



